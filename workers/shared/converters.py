"""Doc-to-text conversion routines.

Each function takes raw bytes + optional filename and returns plain text.
Errors are surfaced to the caller (Temporal will surface them on the
workflow if the activity raises).
"""
from __future__ import annotations
import io
import os
import csv
import json


def _ext(filename: str) -> str:
    return os.path.splitext(filename)[1].lower().lstrip(".")


def convert_docx(data: bytes) -> str:
    from docx import Document  # python-docx
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def convert_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    chunks = []
    for i, page in enumerate(reader.pages):
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        if txt.strip():
            chunks.append(f"[Page {i + 1}]\n{txt}")
    return "\n\n".join(chunks)


def convert_html(data: bytes) -> str:
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(data, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(separator="\n").strip()


def convert_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    return "\n".join("\t".join(row) for row in reader)


def convert_json(data: bytes) -> str:
    try:
        return json.dumps(json.loads(data.decode("utf-8")), indent=2)
    except Exception:
        return data.decode("utf-8", errors="replace")


def convert_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


_DISPATCH = {
    "docx": convert_docx,
    "pdf": convert_pdf,
    "html": convert_html,
    "htm": convert_html,
    "csv": convert_csv,
    "json": convert_json,
    "txt": convert_text,
    "md": convert_text,
    "log": convert_text,
    "py": convert_text,
    "js": convert_text,
    "ts": convert_text,
    "yaml": convert_text,
    "yml": convert_text,
}

# Image extensions that should be routed to the OCR worker instead.
IMAGE_EXTS = {"png", "jpg", "jpeg", "tif", "tiff", "bmp", "gif", "webp"}


def is_image(filename: str, content_type: str = "") -> bool:
    if (content_type or "").startswith("image/"):
        return True
    return _ext(filename) in IMAGE_EXTS


def convert(filename: str, content_type: str, data: bytes) -> str:
    """Pick a converter based on extension first, then content-type."""
    ext = _ext(filename)
    fn = _DISPATCH.get(ext)
    if fn:
        return fn(data)

    ct = (content_type or "").lower()
    if "pdf" in ct:
        return convert_pdf(data)
    if "html" in ct:
        return convert_html(data)
    if "json" in ct:
        return convert_json(data)
    if "csv" in ct:
        return convert_csv(data)
    if "officedocument.wordprocessingml" in ct:
        return convert_docx(data)
    if ct.startswith("text/"):
        return convert_text(data)

    # Last resort: try to decode as text.
    return convert_text(data)
